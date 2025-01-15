from flask import Flask, request, send_file, render_template
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('index.html')

def atualizar_contrato(dados):
    modelo_path = 'contrato_modelo.docx'
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo de contrato não encontrado: {modelo_path}")

    doc = Document(modelo_path)

    substituicoes = {
        'BRENO OLIVEIRA LESSA': dados['nome'],
        '58.851.203/0001-71': dados['cnpj'],
        '58.851.203': dados['cnpj'].split('/')[0],
        'RUA JORDAO 510 – TANQUE - RIO DE JANEIRO - RJ': dados['endereco'],
        '164.866.587-03': dados['cpf'],
        '315429951': dados['rg'],
        'MOTOBOY': dados['tipo_trabalho'],
        'RIO DE JANEIRO, 15 de Janeiro de 2025': dados['data'],
        'X BRENO OLIVEIRA LESSA': dados['assinatura']
    }

    for p in doc.paragraphs:
        for key, value in substituicoes.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in substituicoes.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    output_path = 'contrato_atualizado.docx'
    doc.save(output_path)
    return output_path

@app.route('/gerar_contrato', methods=['POST'])
def gerar_contrato():
    dados = request.json
    dados['data'] = datetime.now().strftime("RIO DE JANEIRO, %d de %B de %Y")
    output_path = atualizar_contrato(dados)
    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)